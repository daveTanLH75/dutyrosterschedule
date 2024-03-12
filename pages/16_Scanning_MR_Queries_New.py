import streamlit as st
import re
import json
import os
import pathlib
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import sqlglot
from sqlglot import parse_one, exp


st.set_page_config(
    page_title="Scanning MR",
    page_icon="👋",
)
st.title("Scanning MR")

def load_sensitive_data_config(config_file):
    #with open(config_file, 'r') as f:
    config = json.load(config_file)
    return {field.lower() for field in config['sensitive_fields']}

def load_approved_sensitive_fields(approved_config_file):
    #with open(approved_config_file, 'r') as f:
    config = json.load(approved_config_file)
    return {field.lower() for field in config['approved_sensitive_fields']}

def categorize_queries(query_results):
    approved_sensitive_queries = []
    unsafe_queries = []
    safe_queries = []
    not_allowed_queries = []

    for filename, query_result, query in query_results:
        if query_result == "APPROVED_SENSITIVE_FIELD_FOUND":
            approved_sensitive_queries.append((filename, query_result, query))
        elif query_result in ["SENSITIVE_FIELD_FOUND", "SELECT_ALL_DATA_FOUND", "SELECT_ALL_FIELDS_FOUND"]:
            unsafe_queries.append((filename, query_result, query))
        elif query_result == "SAFE_QUERY":
            safe_queries.append((filename, query_result, query))
        elif query_result == "NOT_ALLOWED_STATEMENT":
            not_allowed_queries.append((filename, query_result, query))

    return approved_sensitive_queries, unsafe_queries, safe_queries, not_allowed_queries

def validate_queries_in_folder(sensitive_fields, approved_sensitive_fields, queries):
   
    # Get a list of query files in the specified folder
    query_results = []
   
    # Display queries with approved sensitive fields
    for query in queries:
        #st.write(query)
        qur = query[0]
        qur = qur.strip()
        qur_fileName = query[1]
       # st.write(qur)
       # st.write(qur_fileName)
        if qur:
            query_result = is_sensitive_query(qur, sensitive_fields, approved_sensitive_fields)
            query_results.append((qur_fileName, query_result, query))

    approved_sensitive_queries, unsafe_queries, safe_queries, not_allowed_queries = categorize_queries(query_results)

    # Create a new document
    doc = Document()

    # Add a title
    title = doc.add_heading('MR Scripts Scanning Report', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add a table
    doc.add_heading('Queries with Approved Sensitive Fields', level=2)
    tbl_row = len(approved_sensitive_queries) + 1
    table = doc.add_table(rows=tbl_row, cols=3)
    table.style = 'Table Grid'
    table.autofit = False
    table.allow_autofit = False
    for row in table.rows:
        for cell in row.cells:
            cell.width = Pt(100)
    table.cell(0, 0).text = 'Filename'
    table.cell(0, 1).text = 'Query Result'
    table.cell(0, 2).text = 'Query'
    for i, data in enumerate(approved_sensitive_queries, start=1):
        table.cell(i, 0).text = data[0]
        table.cell(i, 1).text = data[1]
        display_txt = data[2][0]
        table.cell(i, 2).text = display_txt
    

    doc.add_heading('Unsafe Queries', level=2)
    tbl_row = len(unsafe_queries) + 1
    table = doc.add_table(rows=tbl_row, cols=3)
    table.style = 'Table Grid'
    table.autofit = False
    table.allow_autofit = False
    for row in table.rows:
        for cell in row.cells:
            cell.width = Pt(100)
    table.cell(0, 0).text = 'Filename'
    table.cell(0, 1).text = 'Query Result'
    table.cell(0, 2).text = 'Query'
    for i, data in enumerate(unsafe_queries, start=1):
        table.cell(i, 0).text = data[0]
        table.cell(i, 1).text = data[1]
        display_txt = data[2][0]
        table.cell(i, 2).text = display_txt

    doc.add_heading('Safe Queries', level=2)
    tbl_row = len(safe_queries) + 1
    table = doc.add_table(rows=tbl_row, cols=3)
    table.style = 'Table Grid'
    table.autofit = False
    table.allow_autofit = False
    for row in table.rows:
        for cell in row.cells:
            cell.width = Pt(100)
    table.cell(0, 0).text = 'Filename'
    table.cell(0, 1).text = 'Query Result'
    table.cell(0, 2).text = 'Query'
    for i, data in enumerate(safe_queries, start=1):
        table.cell(i, 0).text = data[0]
        table.cell(i, 1).text = data[1]
        display_txt = data[2][0]
        table.cell(i, 2).text = display_txt

    doc.add_heading('Not Allowed SQL Statements', level=2)
    tbl_row = len(not_allowed_queries) + 1
    table = doc.add_table(rows=tbl_row, cols=3)
    table.style = 'Table Grid'
    table.autofit = False
    table.allow_autofit = False
    for row in table.rows:
        for cell in row.cells:
            cell.width = Pt(100)
    table.cell(0, 0).text = 'Filename'
    table.cell(0, 1).text = 'Query Result'
    table.cell(0, 2).text = 'Query'
    for i, data in enumerate(not_allowed_queries, start=1):
        table.cell(i, 0).text = data[0]
        table.cell(i, 1).text = data[1]
        display_txt = data[2][0]
        table.cell(i, 2).text = display_txt

    doc.add_heading('Listing of All Queries', level=2)
    tbl_row = len(query_results) + 1
    table = doc.add_table(rows=tbl_row, cols=3)
    table.style = 'Table Grid'
    table.autofit = False
    table.allow_autofit = False
    for row in table.rows:
        for cell in row.cells:
            cell.width = Pt(100)
    table.cell(0, 0).text = 'Filename'
    table.cell(0, 1).text = 'Query Result'
    table.cell(0, 2).text = 'Query'
    for i, data in enumerate(query_results, start=1):
        table.cell(i, 0).text = data[0]
        table.cell(i, 1).text = data[1]
        display_txt = data[2][0]
        table.cell(i, 2).text = display_txt

    doc.save('MR_Scripts_Scan_report.docx')

    bio = io.BytesIO()

    doc.save(bio)
    if doc:
        st.download_button(
            label="Click here to download MR Scanning Report",
            data=bio.getvalue(),
            file_name="MR_Scripts_Scan_report.docx",
            mime="docx"
            )

    # Display queries with approved sensitive fields
    st.text("Output for troubleshoot")
    st.text("\nQueries with Approved Sensitive Fields:")
    st.text(f"{'Filename':<20} {'Query Result':<30} {'Query'}")
    for filename, query_result, query in approved_sensitive_queries:
        display_txt = query[0]
        st.text(f"{filename:<20} {query_result:<30} {display_txt}")

    # Display unsafe queries
    st.text("\nUnsafe Queries:")
    st.text(f"{'Filename':<20} {'Query Result':<30} {'Query'}")
    for filename, query_result, query in unsafe_queries:
        display_txt = query[0]
        st.text(f"{filename:<20} {query_result:<30} {display_txt}")

    # Display safe queries
    st.text("\nSafe Queries:")
    st.text(f"{'Filename':<20} {'Query Result':<30} {'Query'}")
    for filename, query_result, query in safe_queries:
        display_txt = query[0]
        st.text(f"{filename:<20} {query_result:<30} {display_txt}")

    # Display not allowed queries
    st.text("\nNot Allowed SQL Statements:")
    st.text(f"{'Filename':<20} {'Query Result':<30} {'Query'}")
    for filename, query_result, query in not_allowed_queries:
        display_txt = query[0]
        st.text(f"{filename:<20} {query_result:<30} {display_txt}")

    # Display a listing of all queries
    st.text("\nListing of All Queries:")
    st.text(f"{'Filename':<20} {'Query Result':<30} {'Query'}")
    for filename, query_result, query in query_results:
        display_txt = query[0]
        st.text(f"{filename:<20} {query_result:<30} {display_txt}")
    

def is_sensitive_query(query, sensitive_fields, approved_sensitive_fields):
    normalized_query = query.lower()
    normalized_query = normalized_query.strip()

    # Check if the query contains "SELECT *"
    if "select *" in normalized_query:
        return "SELECT_ALL_FIELDS_FOUND"

    # Check if the query contains "SELECT" without specifying columns and lacks a "WHERE" clause
    if "select " in normalized_query and "from " in normalized_query and " where " not in normalized_query:
        return "SELECT_ALL_DATA_FOUND"

    # Check if the query is not a SELECT statement
    if not normalized_query.startswith("select "):
        return "NOT_ALLOWED_STATEMENT"

    # Check if any sensitive field is referenced in the query
    for field in sensitive_fields:
        if field in normalized_query:
            # Check if it's an approved sensitive field
            if field in approved_sensitive_fields:
                return "APPROVED_SENSITIVE_FIELD_FOUND"
            else:
                return "SENSITIVE_FIELD_FOUND"

    # If none of the above conditions are met, the query is safe
    return "SAFE_QUERY"

def initLayout():
    sql_files = []
    sensitive_fields=[]
    approved_sensitive_fields=[]
    queries = []
    FILE_TYPES = ["json","sql"]
    uploaded_files = st.file_uploader("Choose multiple files for scanning (by default it checks for nric, email and home addresses)", type=FILE_TYPES,accept_multiple_files=True)
    for uploaded_file in uploaded_files:
       if uploaded_file.name == 'sensitive_fields.json' :
        sensitive_fields = load_sensitive_data_config(uploaded_file)
       elif uploaded_file.name == 'approved_sensitive_fields.json':
        approved_sensitive_fields = load_approved_sensitive_fields(uploaded_file)
       elif pathlib.Path(uploaded_file.name).suffix == '.sql':
           sql_files.append(uploaded_file)
           data_string = uploaded_file.read()
           str_sql = bytes.decode(data_string)
           str_sql_str = str_sql.split(';')
           for str_sql_str_str in str_sql_str:
               #st.text(str_sql_str_str)
               queries.append((str_sql_str_str,uploaded_file.name))


    if len(sensitive_fields) == 0:
        sensitive_fields = st.secrets["sensitive_fields"]

    if len(approved_sensitive_fields) == 0:
        approved_sensitive_fields = st.secrets["approved_sensitive_fields"]
	
    if len(sensitive_fields) > 0 and len(approved_sensitive_fields) > 0 and len(queries) > 0 :
        st.text("Approved Sensitive Fields:")
        st.write(approved_sensitive_fields)
        st.text("Sensitive Fields:")
        st.write(sensitive_fields)
        validate_queries_in_folder(sensitive_fields,approved_sensitive_fields,queries)
        #st.write(queries[0][1])
        #st.text("Approved Sensitive Fields:")
        #st.write(approved_sensitive_fields)
        #st.text("Sensitive Fields:")
        #st.write(sensitive_fields)
        #st.write(sql_files)

def validateInput():
    sqlstr = st.session_state['sqlQueryScan']
    if sqlstr == "" or sqlstr == None:
        st.error ("No sql queries input",icon="🚨")
        return False
    
    approvedfields = st.session_state['approvedFieldScan']
    if approvedfields == "" or approvedfields == None:
        st.error("No approved fields input",icon="🚨")
        return False


def startScanning():
    if validateInput() == False :
        return
    mrjira_name = ""

    sqlstr = st.session_state['sqlQueryScan']
    approvedfields = st.session_state['approvedFieldScan']
    sqlstrs = sqlstr.split(";")
    appFieldsList = approvedfields.split("\n")
    sensitive_fields = st.secrets["sensitive_fields"]

    #for column in parse_one(sqlstr).find_all(exp.Column):
    #    st.text(f"Column =>  {column.name}")

    #for table in parse_one(sqlstr).find_all(exp.Table):
    #    st.text(f"Table =>  {table.name} | DB=> {table.db}" )

    #st.write(sqlstrs)
    count = 0 # used for first level filtering for jira name and sets of tbl names and queries
    sqlqueries = [] # used to store all the queries, in relation to the table names stored in sqltblNames[]
    sqltblNames = [] # used to store all table names in relation to the queries stored in sqlqueries
    approvedTblNames = [] #used to store all the approved tbl names in relation to the approved fields stored in approvedfieldstrs
    approvedFieldStrs = [] #used to store all the approved fields in relation to approvedtblnames

    tblName = ""
    queryStr = ""
    cnt = 0 # second level filtering for table and queries
    for substr in sqlstrs:
        substrwords = substr.strip().split("\n")
        if len(substrwords) == 0:
            count += 1
            continue

        for substr2 in substrwords:
            substr2 = substr2.strip().lower()
            #st.text(substr2)
            if substr2 == "":
                cnt += 1
                continue

            if "moesyspq" in substr2: #parse jira name
                mrjira_name = substr2
                st.text("Start Scanning MR: "+ mrjira_name)
                cnt += 1
                continue

            if  substr2.startswith("--"): #parse comments happen before qecho
                cnt += 1
                continue

            if "qecho" in substr2: # parse tbl Name
                tblStrlst = substr2.split()
                if len(tblStrlst) == 2:
                    tblName = tblStrlst[1].strip()
                    sqltblNames.append(tblName)
            else:
                if queryStr == "":
                    queryStr = substr2
                else:
                    queryStr = queryStr +" "+ substr2
                    #st.text(queryStr)
            cnt += 1

        #st.text(queryStr)

        if queryStr != "":
            cols = []
            try:
                colls = parse_one(queryStr, dialect="postgres").find_all(exp.Select)
                collscnt = 0
                for column in colls:
                    if collscnt > 0: # this is to skip subqueries inside the main sql. if there are subqueries, colls will have more than 1 count
                        continue
                    for proj in column.expressions:
                        if proj.alias_or_name.lower() not in cols:
                            cols.append(proj.alias_or_name.lower())
                    collscnt += 1
            except sqlglot.errors.ParseError as e:
                    st.error(":red["+queryStr+"] has errors", icon="🚨")
   
            #st.write(cols)
            sqlqueries.append(cols)
            queryStr = ""

        count += 1

    for appfield in appFieldsList:
        
        if "row" in appfield:
            continue
        
        if "+" in appfield:
            continue
        
        if "|" in appfield: #these are fields string
            approvedFieldStrs.append(appfield.strip())
            continue

        if appfield.rstrip() != "":
            approvedTblNames.append(appfield)


    if len(sqltblNames) != len(sqlqueries):
        st.text("There are "+ str(len(sqltblNames))+ " tables and " + str(len(sqlqueries))+ " queries")
        st.error("There are mismatched tables and queries in SQL text scan",icon="🚨")
        return
    
    if len(approvedTblNames) != len(approvedFieldStrs) :
        st.text("There are "+ str(len(approvedTblNames))+ " approved tables and " + str(len(approvedFieldStrs))+ " fields")
        st.error("There are mismatched approved tables and fields in approved fields input", icon="🚨")
        return

    if len(sqltblNames) != len(approvedTblNames):
        st.text("There are "+ str(len(sqltblNames))+ " tables and " + str(len(approvedTblNames))+ " approved tables")
        st.error("There are mismatched tables and approved tables",icon="🚨")
        return
    
    if len(sqlqueries) != len(approvedFieldStrs):
        st.text("There are "+ str(len(sqlqueries))+ " queries and " + str(len(approvedFieldStrs))+ " approved fields")
        st.error("There are mismatched sql queries and approved fields",icon="🚨")
        return

    #start matching
    errorFound = 0
    count = 0
    for actTblName in sqltblNames:
        for appTblName in approvedTblNames:
            if actTblName.strip() == appTblName.strip():
                break
        sqlq = sqlqueries[count] # this is a list now
        appFStr = approvedFieldStrs[count]
        appFStrLst = appFStr.split("|")
        found = False
        for qstr in sqlq:
            found = False
            for aFS in appFStrLst:
                if qstr == aFS.strip():
                    found = True
                    if qstr.strip() != "" and qstr.strip() in sensitive_fields:
                        st.warning("Table: "+ actTblName+ " Field :"+aFS+" is :red[sensitive]", icon="⚠️")
                    break
            if found == False:
                st.text(qstr)
                if qstr.strip() != "" and qstr.strip() in sensitive_fields:
                    st.error("Table Name: "+actTblName+" Field: "+ qstr+" is not matched, :red[This field is sensitve also], Pls check", icon="🚨")
                    errorFound += 1
                elif qstr.strip() != "":
                    st.error("Table Name: "+actTblName+" Field: "+qstr+" is not matched, Pls check", icon="🚨")
                    errorFound += 1
                    
        count += 1
    
    st.text("End of Scanning, "+ str(errorFound)+ " Errors found")

def initLayout2():
    st.header("SQL Queries to Scan")
    st.text_area("Input SQL queries to scan",key="sqlQueryScan", height=500)
    st.header("Approved Fields")
    st.text_area("Input approved database fields to scan",key="approvedFieldScan", height=500)
    st.button("Start Scanning", on_click=startScanning, type= "primary")


initLayout2()
